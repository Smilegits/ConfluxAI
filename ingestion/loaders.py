"""
Source loaders — extract structured text from web pages, PDFs, DOCX, CSV, Excel, and plain text.
Each returns a RawDocument with text, sections, and metadata.
"""
from __future__ import annotations
import csv
import hashlib
import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse

import fitz  # PyMuPDF
import requests
from bs4 import BeautifulSoup
from bs4.element import NavigableString
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass
class RawDocument:
    text: str
    metadata: dict = field(default_factory=dict)
    sections: list[dict] = field(default_factory=list)

    @property
    def content_hash(self) -> str:
        return hashlib.sha256(self.text.encode()).hexdigest()[:16]


class WebLoader:
    STRIP_TAGS = {"nav", "footer", "header", "aside", "script", "style", "noscript", "form"}
    _NAV_RE = re.compile(r'(Previous|Next)\s*$', re.MULTILINE)

    def load(self, url: str) -> RawDocument:
        resp = requests.get(url, timeout=15, headers={"User-Agent": "RAGBot/1.0"})
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup.find_all(self.STRIP_TAGS):
            tag.decompose()

        main = soup.find("main") or soup.find("article") or soup.find("body") or soup
        sections = self._extract_sections(main)
        full_text = "\n\n".join(
            (f"## {s['heading']}\n{s['text']}" if s["heading"] else s["text"]) for s in sections
        )
        if len(full_text.strip()) < 200:
            raise ValueError(
                f"Page returned too little text ({len(full_text.strip())} chars) — "
                "likely requires JavaScript rendering. "
                f"Try ingesting https://r.jina.ai/{url} instead."
            )
        return RawDocument(
            text=full_text,
            metadata={
                "source_type": "web",
                "source": url,
                "title": (soup.title.string.strip() if soup.title and soup.title.string else urlparse(url).netloc),
                "ingested_at": datetime.now(timezone.utc).isoformat(),
            },
            sections=sections,
        )

    def _extract_sections(self, el) -> list[dict]:
        sections, heading, anchor, parts = [], "", "", []
        for child in el.descendants:
            if child.name and child.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                if parts:
                    sections.append({
                        "heading": heading,
                        "anchor": anchor,
                        "text": self._clean(" ".join(parts).strip()),
                    })
                    parts = []
                heading = child.get_text(strip=True)
                anchor = child.get("id", "") or ""
            elif isinstance(child, NavigableString):
                t = child.strip()
                if t and child.parent.name not in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    parts.append(t)
        if parts:
            sections.append({
                "heading": heading,
                "anchor": anchor,
                "text": self._clean(" ".join(parts).strip()),
            })
        return sections

    @classmethod
    def _clean(cls, text: str) -> str:
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)
        text = cls._NAV_RE.sub('', text)
        return text.strip()


class PlaywrightWebLoader(WebLoader):
    """Playwright HTTP-only loader — no browser, works behind corp proxy.
    Activate with USE_PLAYWRIGHT_REQUESTS=true in .env.
    Requires: pip install playwright && playwright install chromium
    Proxy: set HTTP_PROXY / HTTPS_PROXY env vars (Playwright picks them up automatically).
    """

    def load(self, url: str) -> RawDocument:
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            raise ImportError(
                "playwright not installed: pip install playwright && playwright install chromium"
            )

        try:
            import certifi
            os.environ.setdefault("SSL_CERT_FILE", certifi.where())
            os.environ.setdefault("REQUESTS_CA_BUNDLE", certifi.where())
        except ImportError:
            pass  # ignore_https_errors covers corp proxy MITM without certifi

        parsed = urlparse(url)
        base_url = f"{parsed.scheme}://{parsed.netloc}"
        path = url[len(base_url):] or "/"

        with sync_playwright() as p:
            ctx = p.request.new_context(
                base_url=base_url,
                ignore_https_errors=True,
                extra_http_headers={"User-Agent": "RAGBot/1.0"},
            )
            try:
                resp = ctx.get(path)
                html = resp.text()
            finally:
                ctx.dispose()

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup.find_all(self.STRIP_TAGS):
            tag.decompose()

        main = soup.find("main") or soup.find("article") or soup.find("body") or soup
        sections = self._extract_sections(main)
        full_text = "\n\n".join(
            (f"## {s['heading']}\n{s['text']}" if s["heading"] else s["text"]) for s in sections
        )
        if len(full_text.strip()) < 200:
            raise ValueError(
                f"Page returned too little text ({len(full_text.strip())} chars) — "
                "page may require JavaScript rendering. "
                f"Try ingesting https://r.jina.ai/{url} instead."
            )
        return RawDocument(
            text=full_text,
            metadata={
                "source_type": "web",
                "source": url,
                "title": (soup.title.string.strip() if soup.title and soup.title.string else parsed.netloc),
                "ingested_at": datetime.now(timezone.utc).isoformat(),
            },
            sections=sections,
        )


class PDFLoader:
    def load(self, path: str) -> RawDocument:
        doc = fitz.open(path)
        sections, parts = [], []
        for i, page in enumerate(doc, 1):
            text = page.get_text("text").strip()
            if text:
                sections.append({"heading": f"Page {i}", "text": text})
                parts.append(text)
        doc.close()
        return RawDocument(
            text="\n\n".join(parts),
            metadata={"source_type": "pdf", "source": path, "title": Path(path).stem,
                       "pages": len(sections), "ingested_at": datetime.now(timezone.utc).isoformat()},
            sections=sections,
        )


class DocxLoader:
    def load(self, path: str) -> RawDocument:
        doc = DocxDocument(path)
        sections, heading, parts = [], "", []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name.startswith("Heading"):
                if parts:
                    sections.append({"heading": heading, "text": "\n".join(parts)})
                    parts = []
                heading = text
            else:
                parts.append(text)
        if parts:
            sections.append({"heading": heading, "text": "\n".join(parts)})

        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
                sections.append({"heading": f"Table {i+1}", "text": rows[0]+"\n"+sep+"\n"+"\n".join(rows[1:])})

        full = "\n\n".join((f"## {s['heading']}\n{s['text']}" if s["heading"] else s["text"]) for s in sections)
        return RawDocument(
            text=full,
            metadata={"source_type": "docx", "source": path, "title": Path(path).stem,
                       "ingested_at": datetime.now(timezone.utc).isoformat()},
            sections=sections,
        )


class CSVLoader:
    _GROUP_SIZE = 12

    def load(self, path: str) -> RawDocument:
        rows = self._read(path)
        if not rows:
            return RawDocument(
                text="",
                metadata={"source_type": "csv", "source": path, "title": Path(path).stem},
            )

        headers = list(rows[0].keys())
        sections: list[dict] = []

        for i in range(0, len(rows), self._GROUP_SIZE):
            batch = rows[i:i + self._GROUP_SIZE]
            lines = []
            for row in batch:
                parts = [f"{k}: {v}" for k, v in row.items() if str(v).strip()]
                if parts:
                    lines.append(" | ".join(parts))
            if not lines:
                continue
            heading = f"Rows {i + 1}–{min(i + self._GROUP_SIZE, len(rows))}"
            sections.append({"heading": heading, "text": "\n".join(lines)})

        full_text = "\n\n".join(
            (f"## {s['heading']}\n{s['text']}" if s["heading"] else s["text"])
            for s in sections
        )
        return RawDocument(
            text=full_text,
            metadata={
                "source_type": "csv",
                "source": path,
                "title": Path(path).stem,
                "columns": ", ".join(headers),
                "row_count": len(rows),
                "ingested_at": datetime.now(timezone.utc).isoformat(),
            },
            sections=sections,
        )

    @staticmethod
    def _read(path: str) -> list[dict]:
        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                with open(path, newline="", encoding=enc) as f:
                    return list(csv.DictReader(f))
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Cannot decode CSV: {path}")


class ExcelLoader:
    _GROUP_SIZE = 12

    def load(self, path: str) -> RawDocument:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl required for Excel files: pip install openpyxl")

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sections: list[dict] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            raw_rows = [
                [str(cell.value).strip() if cell.value is not None else "" for cell in row]
                for row in ws.iter_rows()
            ]
            raw_rows = [r for r in raw_rows if any(c for c in r)]
            if not raw_rows:
                continue

            # If first row is a title/merged row (≥50% empty cells), skip it
            # and use next row as headers
            def _is_title_row(r: list) -> bool:
                non_empty = sum(1 for c in r if c)
                return len(r) > 1 and non_empty <= max(1, len(r) // 2)

            if _is_title_row(raw_rows[0]) and len(raw_rows) > 1:
                raw_rows = raw_rows[1:]

            headers = raw_rows[0]
            data_rows = raw_rows[1:] if len(raw_rows) > 1 else []
            if not data_rows:
                sections.append({"heading": sheet_name, "text": " | ".join(h for h in headers if h)})
                continue

            for i in range(0, len(data_rows), self._GROUP_SIZE):
                batch = data_rows[i:i + self._GROUP_SIZE]
                lines = []
                for row in batch:
                    if len(headers) == len(row):
                        # Skip pairs where header is empty — use value alone
                        parts = [
                            f"{h}: {v}" if h else v
                            for h, v in zip(headers, row) if v
                        ]
                    else:
                        parts = [v for v in row if v]
                    if parts:
                        lines.append(" | ".join(parts))
                if not lines:
                    continue
                end_row = min(i + self._GROUP_SIZE, len(data_rows))
                heading = f"{sheet_name} — rows {i + 2}–{end_row + 1}"
                sections.append({"heading": heading, "text": "\n".join(lines)})

        sheet_count = len(wb.sheetnames)  # read before close
        wb.close()

        full_text = "\n\n".join(
            (f"## {s['heading']}\n{s['text']}" if s["heading"] else s["text"])
            for s in sections
        )
        return RawDocument(
            text=full_text,
            metadata={
                "source_type": "excel",
                "source": path,
                "title": Path(path).stem,
                "sheets": sheet_count,
                "ingested_at": datetime.now(timezone.utc).isoformat(),
            },
            sections=sections,
        )


class TxtLoader:
    def load(self, path: str) -> RawDocument:
        text = Path(path).read_text(encoding="utf-8", errors="replace")
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        sections = [{"heading": "", "text": p} for p in paragraphs]
        return RawDocument(
            text=text,
            metadata={
                "source_type": "text",
                "source": path,
                "title": Path(path).stem,
                "ingested_at": datetime.now(timezone.utc).isoformat(),
            },
            sections=sections,
        )


class TextLoader:
    """For raw pasted text."""
    def load(self, text: str, title: str = "Pasted Text") -> RawDocument:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        sections = [{"heading": "", "text": p} for p in paragraphs]
        return RawDocument(
            text=text,
            metadata={"source_type": "text", "source": "manual", "title": title,
                       "ingested_at": datetime.now(timezone.utc).isoformat()},
            sections=sections,
        )


def get_loader(source: str):
    s = source.lower()
    if source.startswith(("http://", "https://")):
        from config import settings
        if settings.use_playwright_requests:
            return PlaywrightWebLoader(), "url"
        return WebLoader(), "url"
    elif s.endswith(".pdf"):
        return PDFLoader(), "file"
    elif s.endswith(".docx") or s.endswith(".doc"):
        return DocxLoader(), "file"
    elif s.endswith(".csv"):
        return CSVLoader(), "file"
    elif s.endswith(".xlsx") or s.endswith(".xls"):
        return ExcelLoader(), "file"
    elif s.endswith(".txt"):
        return TxtLoader(), "file"
    else:
        raise ValueError(f"Unsupported: {source}. Supported: URL, .pdf, .docx, .csv, .xlsx, .xls, .txt")
